from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "reverse-engineering" / "05-use-cases.md"
OUTPUT = ROOT / "Ma_tran_phan_quyen_he_thong.docx"


def set_font(run, size=9, bold=False, color="1F2937"):
    run.font.name = "Calibri"
    rpr = run._element.get_or_add_rPr()
    for key in ("ascii", "hAnsi", "eastAsia"):
        rpr.rFonts.set(qn(f"w:{key}"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    node = tcpr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        tcpr.append(node)
    node.set(qn("w:fill"), fill)


def margins(cell, top=80, start=90, bottom=80, end=90):
    tcpr = cell._tc.get_or_add_tcPr()
    mar = tcpr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tcpr.append(mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_width(cell, width):
    tcpr = cell._tc.get_or_add_tcPr()
    tcw = tcpr.find(qn("w:tcW"))
    if tcw is None:
        tcw = OxmlElement("w:tcW")
        tcpr.append(tcw)
    tcw.set(qn("w:w"), str(width))
    tcw.set(qn("w:type"), "dxa")


def geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblpr = table._tbl.tblPr
    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.append(tblw)
    tblw.set(qn("w:w"), str(sum(widths)))
    tblw.set(qn("w:type"), "dxa")
    tblind = tblpr.find(qn("w:tblInd"))
    if tblind is None:
        tblind = OxmlElement("w:tblInd")
        tblpr.append(tblind)
    tblind.set(qn("w:w"), "90")
    tblind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_width(cell, width)
            cell.width = Inches(width / 1440)


def row_options(row, header=False):
    trpr = row._tr.get_or_add_trPr()
    no_split = OxmlElement("w:cantSplit")
    trpr.append(no_split)
    if header:
        repeat = OxmlElement("w:tblHeader")
        repeat.set(qn("w:val"), "true")
        trpr.append(repeat)


def page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Trang ")
    set_font(run, size=8.5, color="6B7280")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def usecase_names():
    text = SOURCE.read_text(encoding="utf-8")
    found = re.findall(r"^### UC-(\d+) — (.+)$", text, re.MULTILINE)
    rows = [(int(number), name.strip()) for number, name in found]
    if len(rows) != 53:
        raise RuntimeError(f"Expected 53 use cases, found {len(rows)}")
    return rows


PUBLIC = "✓ Công khai"
OWN = "✓ Cá nhân"
YES = "✓"
NO = "—"
ASSIGNED = "✓ Phân công"


def permissions(number):
    # Column order: admin, tour guide, support staff, customer.
    if number in range(1, 9):
        return [PUBLIC, PUBLIC, PUBLIC, PUBLIC]
    if number == 9:
        return [OWN, OWN, OWN, OWN]
    if number in range(10, 19):
        return [NO, NO, NO, OWN]
    if number in range(19, 43):
        return [YES, NO, NO, NO]
    if number in (43, 44, 49, 50):
        return [NO, OWN, NO, NO]
    if number in (45, 46, 47, 48):
        return [NO, ASSIGNED, NO, NO]
    if number in (51, 53):
        return [NO, NO, OWN, NO]
    if number == 52:
        return [NO, NO, YES, NO]
    raise RuntimeError(f"No permission mapping for UC-{number:03d}")


def build():
    rows = usecase_names()
    doc = Document()
    section = doc.sections[0]
    # Named override for this six-column matrix: Letter landscape, 0.5-inch side margins.
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(9)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.05

    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(hp.add_run("MA TRẬN PHÂN QUYỀN HỆ THỐNG"), size=8, bold=True, color="6B7280")
    page_number(section.footer.paragraphs[0])

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(3)
    set_font(title.add_run("MA TRẬN PHÂN QUYỀN CỦA HỆ THỐNG"), size=17, bold=True, color="1F4D78")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(5)
    set_font(
        subtitle.add_run("Đối chiếu theo middleware, vai trò, quyền sở hữu dữ liệu và phạm vi tour được phân công"),
        size=9,
        color="5B6573",
    )

    legend = doc.add_paragraph()
    legend.alignment = WD_ALIGN_PARAGRAPH.CENTER
    legend.paragraph_format.space_after = Pt(8)
    set_font(legend.add_run("✓: Có quyền   |   Cá nhân: Chỉ dữ liệu của mình   |   Phân công: Chỉ tour được giao   |   —: Không có quyền"), size=8.5, bold=True, color="334155")

    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    widths = [600, 4200, 2400, 2400, 2400, 2400]
    labels = ["STT", "Chức năng", "Quản trị viên", "Hướng dẫn viên", "Nhân viên hỗ trợ", "Khách hàng"]
    for cell, label in zip(table.rows[0].cells, labels):
        shade(cell, "D9EAF7")
        margins(cell, top=105, bottom=105)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(label), size=9, bold=True, color="163A5F")
    row_options(table.rows[0], header=True)

    for number, function in rows:
        perms = permissions(number)
        cells = table.add_row().cells
        values = [str(number), function, *perms]
        for index, (cell, value) in enumerate(zip(cells, values)):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            margins(cell)
            if number % 2 == 0:
                shade(cell, "F7FAFC")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if index == 1 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            color = "166534" if value.startswith("✓") else "64748B" if value == NO else "1F2937"
            set_font(p.add_run(value), size=8.6, bold=(index == 1 or value.startswith("✓")), color=color)
        row_options(table.rows[-1])

    geometry(table, widths)

    source_note = doc.add_paragraph()
    source_note.paragraph_format.space_before = Pt(6)
    set_font(source_note.add_run("Ghi chú: "), size=8, bold=True, color="5B6573")
    set_font(source_note.add_run("Ma trận phản ánh quyền backend quan sát được trong source; một số API có thể chưa có giao diện React hoàn chỉnh."), size=8, color="5B6573")

    props = doc.core_properties
    props.title = "Ma trận phân quyền của hệ thống"
    props.subject = "Phân quyền Quản trị viên, Hướng dẫn viên, Nhân viên hỗ trợ và Khách hàng"
    props.author = "Đồ án DATN_WD-01"
    props.keywords = "ma trận phân quyền, RBAC, Use Case"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
