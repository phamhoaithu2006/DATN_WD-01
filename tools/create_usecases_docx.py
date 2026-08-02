from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "reverse-engineering" / "05-use-cases.md"
OUTPUT = ROOT / "Bang_phan_tich_UseCase_he_thong.docx"


def set_font(run, name="Calibri", size=10, bold=False, color="000000"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            cell.width = Inches(width / 1440)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Trang ")
    set_font(run, size=9, color="666666")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def extract_usecases():
    text = SOURCE.read_text(encoding="utf-8")
    pattern = re.compile(
        r"^### UC-(\d+) — (.+?)\n(?:.*?\n)*?- \*\*Mục tiêu[^:]*:\*\* (.+?)$",
        re.MULTILINE,
    )
    rows = []
    for number, name, goal in pattern.findall(text):
        clean_goal = re.sub(r"`([^`]+)`", r"\1", goal).strip()
        rows.append((int(number), name.strip(), clean_goal))
    if len(rows) != 53:
        raise RuntimeError(f"Expected 53 use cases, found {len(rows)}")
    return rows


def build():
    rows = extract_usecases()
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.08

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = header.add_run("PHÂN TÍCH HỆ THỐNG DU LỊCH")
    set_font(hr, size=8.5, bold=True, color="6B7280")
    add_page_number(section.footer.paragraphs[0])

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("PHÂN TÍCH CÁC USE CASE CỦA HỆ THỐNG")
    set_font(tr, size=17, bold=True, color="1F4D78")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("Tổng hợp 53 chức năng từ mã nguồn backend, frontend và tài liệu reverse-engineering")
    set_font(sr, size=9.5, color="5B6573")

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    widths = [650, 2600, 6110]
    headers = ["STT", "Tên UseCase", "Ý nghĩa/Ghi chú"]
    for idx, (cell, label) in enumerate(zip(table.rows[0].cells, headers)):
        shade_cell(cell, "D9EAF7")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell, top=110, bottom=110)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(label)
        set_font(run, size=10, bold=True, color="163A5F")
    repeat_header(table.rows[0])

    for stt, name, note in rows:
        cells = table.add_row().cells
        values = [str(stt), name, note]
        for col, (cell, value) in enumerate(zip(cells, values)):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if stt % 2 == 0:
                shade_cell(cell, "F7FAFC")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col == 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            run = p.add_run(value)
            set_font(run, size=9.3, bold=(col == 1), color="1F2937")
        prevent_row_split(table.rows[-1])

    set_table_geometry(table, widths)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(8)
    note.paragraph_format.space_after = Pt(0)
    nr = note.add_run("Nguồn đối chiếu: ")
    set_font(nr, size=8.5, bold=True, color="5B6573")
    nr2 = note.add_run("docs/reverse-engineering/05-use-cases.md và các route của backend/frontend.")
    set_font(nr2, size=8.5, color="5B6573")

    props = doc.core_properties
    props.title = "Phân tích các Use Case của hệ thống"
    props.subject = "Bảng tổng hợp 53 Use Case"
    props.author = "Đồ án DATN_WD-01"
    props.keywords = "Use Case, phân tích hệ thống, du lịch"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
