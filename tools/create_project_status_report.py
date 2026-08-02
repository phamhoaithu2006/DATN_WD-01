from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parents[1] / "docs" / "Bao_cao_tien_do_du_an_ViVuGo.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
TEXT = RGBColor(0x20, 0x20, 0x20)
MUTED = RGBColor(0x66, 0x66, 0x66)


def set_run_font(run, size=11, bold=None, color=TEXT, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margin(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tcMar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def dont_split_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    elem = OxmlElement("w:cantSplit")
    tr_pr.append(elem)


def set_table_geometry(table, widths_dxa):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths_dxa):
        col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            cell.width = Inches(width / 1440)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margin(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def write_cell(cell, value, *, header=False, center=False, font_size=9.2):
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(str(value))
    set_run_font(run, size=font_size, bold=header, color=RGBColor(0, 0, 0) if header else TEXT)


def add_table(doc, headers, rows, widths, status_col=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for index, value in enumerate(headers):
        set_cell_shading(hdr.cells[index], LIGHT_BLUE)
        write_cell(hdr.cells[index], value, header=True, center=index in (0, len(headers) - 1), font_size=9.3)
    for data in rows:
        cells = table.add_row().cells
        dont_split_row(table.rows[-1])
        for index, value in enumerate(data):
            if status_col == index and "Chưa" in str(value):
                set_cell_shading(cells[index], "FFF2CC")
            elif status_col == index and "một phần" in str(value).lower():
                set_cell_shading(cells[index], "FFF2CC")
            write_cell(cells[index], value, center=index == 0 or index == len(data) - 1, font_size=8.8)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.10
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    run = p.add_run(text)
    set_run_font(run, size=11)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    return p


def configure_document(doc):
    section = doc.sections[0]
    section.orientation = WD_SECTION.NEW_PAGE if False else section.orientation
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run("Báo cáo tiến độ dự án ViVuGo")
    set_run_font(run, size=8.5, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_after = Pt(0)
    run = footer.add_run("Tài liệu tổng hợp từ mã nguồn và kiểm thử hiện trạng")
    set_run_font(run, size=8.5, color=MUTED)


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(16)
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("BÁO CÁO TIẾN ĐỘ DỰ ÁN VIVUGO")
    set_run_font(run, size=22, bold=True, color=RGBColor.from_string(DARK_BLUE))
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    run = subtitle.add_run("Tổng hợp mức độ hoàn thành, rủi ro và kế hoạch phát triển")
    set_run_font(run, size=11.5, color=MUTED, italic=True)

    add_heading(doc, "2. Mức độ hoàn thành dự án")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Đánh giá tổng quan: ")
    set_run_font(r, bold=True)
    r = p.add_run("Hệ thống ViVuGo đã hoàn thành phần lớn luồng nghiệp vụ cốt lõi, ước tính 80-85% theo phạm vi mã nguồn hiện có. Frontend build production thành công; bộ kiểm thử backend hiện chưa thể chạy đầy đủ do môi trường PHP thiếu driver pdo_sqlite.")
    set_run_font(r)

    add_heading(doc, "Bảng trạng thái các chức năng của hệ thống", level=2)
    functions = [
        (1, "Đăng ký, đăng nhập, đăng xuất, quên/đặt lại mật khẩu OTP", "Hoàn thành"),
        (2, "Phân quyền Customer, Admin, Hướng dẫn viên, Nhân viên hỗ trợ", "Hoàn thành"),
        (3, "Trang chủ, danh mục, tìm kiếm và xem chi tiết tour", "Hoàn thành"),
        (4, "Wishlist/yêu thích tour", "Hoàn thành"),
        (5, "Đặt tour, hành khách, giữ chỗ và lịch sử booking", "Hoàn thành"),
        (6, "Thanh toán VNPay, callback/IPN, tiếp tục hoặc hủy thanh toán", "Hoàn thành"),
        (7, "Quản trị tour, lịch khởi hành, giá, hành trình và hình ảnh", "Hoàn thành"),
        (8, "Quản lý khách hàng, hướng dẫn viên và nhân viên hỗ trợ", "Hoàn thành"),
        (9, "Phân công/thay thế hướng dẫn viên, xin nghỉ", "Hoàn thành"),
        (10, "Điểm danh khách và theo dõi tiến độ tour", "Hoàn thành"),
        (11, "Yêu cầu hỗ trợ, trao đổi và xử lý ticket", "Hoàn thành"),
        (12, "Thông báo theo vai trò và chiến dịch thông báo", "Hoàn thành"),
        (13, "Đánh giá tour và hướng dẫn viên; kiểm duyệt đánh giá", "Hoàn thành"),
        (14, "Báo cáo, thống kê và xuất CSV phía trình duyệt", "Hoàn thành"),
        (15, "Cấu hình hệ thống, banner/widget, ngôn ngữ, chứng chỉ", "Hoàn thành"),
        (16, "Sao lưu cơ sở dữ liệu", "Backend hoàn thành; UI một phần"),
        (17, "Chatbot/trợ lý tư vấn tour AI", "Hoàn thành cơ bản"),
        (18, "Khuyến mãi, hoàn tiền, blog, đối tác dịch vụ", "Chưa hoàn thiện"),
        (19, "Xác thực hai lớp (2FA)", "Chưa hoàn thiện"),
        (20, "Tối ưu hiệu năng frontend", "Cần cải thiện"),
    ]
    add_table(doc, ["STT", "Chức năng", "Trạng thái"], functions, [600, 5900, 2860], status_col=2)

    add_heading(doc, "3. Những khó khăn, rủi ro gặp phải và cách giải quyết")
    add_heading(doc, "3.1. Khó khăn, rủi ro gặp phải", level=2)
    risks = [
        ("Cạnh tranh dữ liệu ở luồng đặt tour, thanh toán, hủy booking và phân công HDV", "Có thể giữ chỗ sai, hoàn chỗ nhiều lần hoặc ghi đè trạng thái", "Đã áp dụng transaction, khóa bản ghi và kiểm tra lại trạng thái; duy trì test đồng thời trên MySQL."),
        ("Đồng bộ trạng thái booking và payment", "Booking đã hủy nhưng payment vẫn có thể cập nhật không nhất quán", "Giới hạn các chuyển trạng thái hợp lệ; cập nhật booking/payment trong cùng transaction."),
        ("Bảo mật OTP đặt lại mật khẩu", "Lộ OTP hoặc sử dụng mã OTP quá hạn", "Đã băm OTP, gửi qua email, đặt hạn dùng 10 phút và xóa OTP sau khi sử dụng."),
        ("Môi trường kiểm thử backend thiếu pdo_sqlite", "Không thể xác minh toàn bộ regression test", "Cài/bật extension pdo_sqlite, chạy lại php artisan test và cấu hình CI chạy SQLite/MySQL."),
        ("2FA mới dừng ở mức cấu hình", "Thiết lập bảo mật hai lớp chưa được áp dụng thực tế", "Bổ sung luồng OTP/app authenticator, middleware kiểm tra 2FA và màn hình xác minh."),
        ("Một số module chưa end-to-end: hoàn tiền, promotion, blog, partner", "Giảm tính đầy đủ khi đưa vào vận hành thực tế", "Xác nhận yêu cầu nghiệp vụ; bổ sung workflow, API, UI, phân quyền và test."),
        ("Backup chưa có luồng khôi phục dữ liệu hoàn chỉnh", "Khó phục hồi khi xảy ra sự cố", "Bổ sung UI vận hành backup, quy trình restore, phân quyền và chỉ tiêu RPO/RTO."),
        ("Bundle frontend lớn", "Tăng thời gian tải trang đầu tiên", "Lazy-load các trang quản trị/báo cáo và chia nhỏ bundle bằng dynamic import."),
    ]
    add_table(doc, ["Khó khăn/rủi ro", "Ảnh hưởng", "Cách xử lý"], risks, [3000, 2700, 3660])

    add_heading(doc, "3.2. Cách giải quyết", level=2)
    solutions = [
        "Hoàn thiện môi trường kiểm thử bằng cách cài và kích hoạt pdo_sqlite; chạy lại toàn bộ php artisan test để xác minh chức năng backend.",
        "Duy trì transaction, khóa bản ghi và kiểm tra trạng thái trước khi cập nhật ở các luồng đặt tour, thanh toán, hủy booking, xin nghỉ và phân công hướng dẫn viên.",
        "Hoàn thiện 2FA bằng OTP hoặc ứng dụng xác thực; chỉ cấp token sau khi xác minh thành công.",
        "Xây dựng đầy đủ quy trình hoàn tiền: tạo yêu cầu, duyệt/từ chối, cập nhật payment, lưu lịch sử và gửi thông báo.",
        "Hoàn thiện module khuyến mãi, blog và đối tác dịch vụ với dữ liệu, API, giao diện, phân quyền và kiểm thử tương ứng.",
        "Bổ sung chức năng khôi phục dữ liệu từ bản sao lưu; quy định quyền backup/restore và thời gian khôi phục dữ liệu.",
        "Tối ưu hiệu năng frontend bằng lazy loading, code splitting và tách các trang báo cáo/quản trị thành bundle riêng.",
        "Xây dựng tài liệu KPI, SLA, quy trình xử lý sự cố và phân công trách nhiệm cho từng vai trò vận hành.",
    ]
    for item in solutions:
        add_bullet(doc, item)

    add_heading(doc, "4. Kế hoạch phát triển trong tương lai")
    phases = [
        ("Giai đoạn 1: Hoàn thiện và ổn định hệ thống", [
            "Cài đặt đầy đủ môi trường kiểm thử, chạy toàn bộ test backend và khắc phục lỗi phát sinh.",
            "Hoàn thiện 2FA, quy trình hoàn tiền, mã giảm giá, quản lý đối tác và blog.",
            "Bổ sung chức năng khôi phục dữ liệu và chuẩn hóa luồng trạng thái nghiệp vụ.",
        ]),
        ("Giai đoạn 2: Tối ưu trải nghiệm và hiệu năng", [
            "Tối ưu giao diện di động, chia nhỏ bundle frontend và lazy-load các trang nặng.",
            "Hoàn thiện các màn hình quản trị backup, khuyến mãi và báo cáo chuyên sâu.",
            "Cải thiện chatbot tư vấn tour và mở rộng hỗ trợ đa ngôn ngữ.",
        ]),
        ("Giai đoạn 3: Mở rộng nghiệp vụ", [
            "Tích hợp thêm cổng thanh toán và hóa đơn điện tử.",
            "Phát triển ứng dụng di động cho khách hàng và hướng dẫn viên.",
            "Tích hợp SMS, email marketing, CRM và dashboard phân tích kinh doanh.",
        ]),
        ("Giai đoạn 4: Vận hành và triển khai thực tế", [
            "Triển khai CI/CD, tự động kiểm thử và kiểm tra chất lượng mã nguồn.",
            "Thiết lập giám sát, sao lưu định kỳ, cảnh báo lỗi và nhật ký hoạt động.",
            "Xây dựng KPI, SLA; thực hiện kiểm thử bảo mật, kiểm thử tải và đánh giá khả năng mở rộng.",
        ]),
    ]
    for heading, items in phases:
        add_heading(doc, heading, level=2)
        for item in items:
            add_bullet(doc, item)

    doc.core_properties.title = "Báo cáo tiến độ dự án ViVuGo"
    doc.core_properties.subject = "Mức độ hoàn thành, rủi ro và kế hoạch phát triển"
    doc.core_properties.author = "ViVuGo"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
