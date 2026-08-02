from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parents[1] / "docs" / "Khao_sat_he_thong_ViVuGo.docx"
BLUE = "2E74B5"
DARK = "1F4D78"
PALE = "E8EEF5"
TEXT = RGBColor(32, 32, 32)
MUTED = RGBColor(95, 95, 95)


def font(run, size=11, bold=None, color=TEXT, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def cell_margin(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side, amount in (("top", 80), ("start", 120), ("bottom", 80), ("end", 120)):
        item = OxmlElement(f"w:{side}")
        item.set(qn("w:w"), str(amount))
        item.set(qn("w:type"), "dxa")
        mar.append(item)
    tc_pr.append(mar)


def table_geometry(table, widths):
    table.autofit = False
    pr = table._tbl.tblPr
    tw = pr.first_child_found_in("w:tblW")
    if tw is None:
        tw = OxmlElement("w:tblW")
        pr.append(tw)
    tw.set(qn("w:w"), str(sum(widths)))
    tw.set(qn("w:type"), "dxa")
    ind = OxmlElement("w:tblInd")
    ind.set(qn("w:w"), "120")
    ind.set(qn("w:type"), "dxa")
    pr.append(ind)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    pr.append(layout)
    for col, width in zip(table._tbl.tblGrid.gridCol_lst, widths):
        col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tcw = tc_pr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tc_pr.append(tcw)
            tcw.set(qn("w:w"), str(width))
            tcw.set(qn("w:type"), "dxa")
            cell_margin(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def write_cell(cell, text, header=False, center=False):
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(text)
    font(r, 9.3 if header else 9.1, bold=header, color=RGBColor(0, 0, 0) if header else TEXT)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, head in enumerate(headers):
        shade(table.rows[0].cells[i], PALE)
        write_cell(table.rows[0].cells[i], head, header=True, center=i == 0)
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            write_cell(cells[i], value, center=i == 0)
    table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(.5)
    p.paragraph_format.first_line_indent = Inches(-.25)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    font(p.add_run(text))


def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    font(p.add_run(text), 16 if level == 1 else 13, bold=True, color=RGBColor.from_string(BLUE))


def body(doc, text, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    font(p.add_run(text), italic=italic)


def configure(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = Inches(.492)
    section.footer_distance = Inches(.492)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, before, after in (("Heading 1", 16, 16, 8), ("Heading 2", 13, 12, 6)):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(BLUE)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
    h = section.header.paragraphs[0]
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font(h.add_run("Khảo sát hệ thống ViVuGo"), 8.5, color=MUTED)
    f = section.footer.paragraphs[0]
    f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(f.add_run("Tài liệu khảo sát - 08/2026"), 8.5, color=MUTED)


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    font(p.add_run("KHẢO SÁT HỆ THỐNG"), 22, bold=True, color=RGBColor.from_string(DARK))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    font(p.add_run("Hệ thống quản lý và đặt tour du lịch trực tuyến ViVuGo"), 11.5, italic=True, color=MUTED)

    heading(doc, "2.1. Bài toán nghiệp vụ")
    body(doc, "ViVuGo được xây dựng để số hóa quy trình cung cấp và vận hành tour du lịch: từ công bố tour, tìm kiếm, đặt chỗ và thanh toán của khách hàng đến quản trị tour, lịch khởi hành, hướng dẫn viên, hỗ trợ khách hàng và báo cáo của doanh nghiệp.")
    body(doc, "Thực tế vận hành du lịch có nhiều dữ liệu và vai trò liên quan. Nếu quản lý thủ công hoặc bằng công cụ rời rạc, doanh nghiệp dễ gặp tình trạng chậm xác nhận chỗ, sai lệch số lượng khách, khó theo dõi thanh toán, phân công hướng dẫn viên thiếu đồng bộ và phản hồi hỗ trợ không kịp thời.")
    heading(doc, "Các vấn đề cần giải quyết", level=2)
    for item in [
        "Tập trung thông tin tour, điểm đến, lịch khởi hành, giá và số chỗ còn lại trên một nền tảng duy nhất.",
        "Cho phép khách hàng tìm kiếm, xem chi tiết tour, lưu yêu thích, đặt tour, khai báo hành khách và thanh toán trực tuyến qua VNPay.",
        "Đảm bảo số chỗ, booking và payment được cập nhật nhất quán; hạn chế trùng lặp hoặc thay đổi trạng thái không hợp lệ.",
        "Hỗ trợ bộ phận quản trị quản lý tour, lịch khởi hành, khách hàng, thanh toán, thông báo, báo cáo và cấu hình hệ thống.",
        "Hỗ trợ hướng dẫn viên xem tour được phân công, điểm danh khách, cập nhật tiến độ, gửi yêu cầu thay thế hoặc xin nghỉ.",
        "Tạo kênh hỗ trợ để khách gửi yêu cầu, nhân viên tiếp nhận/xử lý và theo dõi trạng thái yêu cầu.",
    ]:
        bullet(doc, item)
    heading(doc, "Mô hình nghiệp vụ đề xuất", level=2)
    add_table(doc, ["Đối tượng", "Nhu cầu chính", "Chức năng ViVuGo đáp ứng"], [
        ("Khách hàng", "Tìm và đặt tour thuận tiện, minh bạch", "Catalog, tìm kiếm/lọc, chi tiết tour, wishlist, booking, thanh toán, đánh giá, hỗ trợ"),
        ("Quản trị viên", "Điều hành và kiểm soát dữ liệu", "Quản lý tour/lịch/booking/payment, người dùng, báo cáo, thông báo, cấu hình"),
        ("Hướng dẫn viên", "Thực hiện tour và phối hợp vận hành", "Tour được phân công, danh sách khách, điểm danh, tiến độ, xin nghỉ/thay thế"),
        ("Nhân viên hỗ trợ", "Xử lý vấn đề của khách nhanh chóng", "Tiếp nhận, cập nhật trạng thái yêu cầu hỗ trợ và gửi thông báo"),
    ], [1600, 3200, 4560])

    heading(doc, "2.2. Hệ thống tương tự")
    body(doc, "Trong quá trình khảo sát, nhóm chúng em đã tham khảo một số website du lịch uy tín nhằm học hỏi mô hình hoạt động và giao diện quản lý. Hai ví dụ tiêu biểu là Traveloka và Vietravel.")
    heading(doc, "1. Traveloka", level=2)
    body(doc, "Traveloka là nền tảng du lịch trực tuyến đa dịch vụ, cung cấp đặt chỗ lưu trú, vé máy bay, xe đưa đón, thuê xe, hoạt động/trải nghiệm và một số tiện ích đi kèm. Website có cơ chế tìm kiếm theo điểm đến, ngày đi, số khách; hiển thị ưu đãi, đánh giá và hỗ trợ quản lý sau đặt chỗ.")
    body(doc, "Điểm tham khảo cho ViVuGo: bố cục tìm kiếm nổi bật, bộ lọc rõ ràng, trang chi tiết minh bạch, thông báo tức thời và khu vực quản lý booking sau thanh toán.", italic=True)
    heading(doc, "2. Vietravel", level=2)
    body(doc, "Vietravel là website bán tour trực tuyến, tập trung vào tour trong nước và quốc tế, đồng thời cung cấp thêm vé máy bay, khách sạn và các chương trình ưu đãi. Cách phân loại tour theo điểm đến, thời gian, mùa du lịch và chủ đề giúp khách hàng dễ lựa chọn sản phẩm phù hợp.")
    body(doc, "Điểm tham khảo cho ViVuGo: cách trình bày tour trọn gói, lịch khởi hành, nội dung chương trình tour, giá bán và các nhóm tour theo nhu cầu khách hàng Việt Nam.", italic=True)
    heading(doc, "Ưu điểm có thể học hỏi", level=2)
    for item in [
        "Tìm kiếm và lọc tour trực quan theo điểm đến, ngày khởi hành, mức giá và nhu cầu du lịch.",
        "Trang chi tiết tour đầy đủ thông tin: lịch trình, dịch vụ bao gồm/không bao gồm, giá, điều kiện và lịch khởi hành.",
        "Tăng sự tin cậy bằng đánh giá, ưu đãi rõ ràng, xác nhận booking và hỗ trợ sau khi đặt.",
        "Tổ chức danh mục theo chủ đề, mùa và điểm đến để khách khám phá sản phẩm nhanh hơn.",
        "Thiết kế giao diện ưu tiên thao tác đặt dịch vụ, phù hợp cả web và thiết bị di động.",
    ]:
        bullet(doc, item)
    heading(doc, "Nhược điểm và điểm cần lưu ý", level=2)
    for item in [
        "Mô hình đa dịch vụ như Traveloka có phạm vi rất lớn, đòi hỏi dữ liệu, đối tác, vận hành và chi phí tích hợp cao; không phù hợp để triển khai toàn bộ ngay trong giai đoạn đầu.",
        "Nhiều chương trình khuyến mãi hoặc tùy chọn sản phẩm có thể làm giao diện phức tạp, gây khó khăn cho người dùng mới.",
        "Website bán tour có lượng nội dung lớn cần được quản trị thường xuyên; thông tin lịch khởi hành, giá và chính sách phải luôn chính xác.",
        "ViVuGo cần ưu tiên độ tin cậy của booking, thanh toán và vận hành tour trước khi mở rộng sang nhiều dịch vụ phụ trợ.",
    ]:
        bullet(doc, item)

    heading(doc, "Kết luận khảo sát", level=2)
    body(doc, "ViVuGo định hướng phù hợp với mô hình quản lý và đặt tour trực tuyến. Hệ thống nên tiếp thu cách tổ chức tìm kiếm, hiển thị tour, quản lý booking và chăm sóc khách hàng từ Traveloka; đồng thời học hỏi cách phân loại và trình bày tour trọn gói phù hợp thị trường Việt Nam từ Vietravel. Trong giai đoạn đầu, ưu tiên cần đặt vào quy trình tour - booking - thanh toán - vận hành hướng dẫn viên - hỗ trợ khách hàng để tạo nền tảng ổn định.")

    heading(doc, "Nguồn tham khảo", level=2)
    for source in [
        "Tài liệu nội bộ: docs/reverse-engineering/01-executive-domain-analysis.md và 02-module-analysis.md (mã nguồn ViVuGo).",
        "Traveloka, trang giới thiệu và trang chủ: https://www.traveloka.com/en-vn/about-us ; https://www.traveloka.com/en-vn (truy cập tháng 08/2026).",
        "Vietravel, website đặt tour trực tuyến: https://travel.com.vn/ ; trang giới thiệu doanh nghiệp: https://www.vietravel.com/vn/tin-tuc-du-lich/gioi-thieu-ve-cong-ty-vietravel-v14365.aspx (truy cập tháng 08/2026).",
    ]:
        bullet(doc, source)

    doc.core_properties.title = "Khảo sát hệ thống ViVuGo"
    doc.core_properties.author = "Nhóm phát triển ViVuGo"
    doc.core_properties.subject = "Bài toán nghiệp vụ và hệ thống tương tự"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
