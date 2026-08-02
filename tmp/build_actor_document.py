from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path


OUT = Path(r"C:\Users\ADMIN\DATN_WD-01\docs\Tac_nhan_he_thong_ViVuGo.docx")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def write_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(10.5)
    if color:
        run.font.color.rgb = RGBColor(*color)


def write_bullets(cell, items):
    cell.text = ""
    for index, item in enumerate(items):
        p = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        p.style = "List Bullet"
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.05
        run = p.add_run(item)
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        run.font.size = Pt(10)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

normal = doc.styles["Normal"]
normal.font.name = "Arial"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.15

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(4)
run = title.add_run("CÁC TÁC NHÂN CỦA HỆ THỐNG")
run.bold = True
run.font.name = "Arial"
run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(31, 78, 121)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(16)
subrun = subtitle.add_run("Hệ thống quản lý và đặt tour du lịch ViVuGo")
subrun.italic = True
subrun.font.name = "Arial"
subrun._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
subrun.font.size = Pt(11)
subrun.font.color.rgb = RGBColor(89, 89, 89)

intro = doc.add_paragraph()
intro.paragraph_format.space_after = Pt(10)
intro.add_run("Hệ thống có 04 tác nhân chính. Danh sách dưới đây tổng hợp đầy đủ các chức năng đã được xác minh trong mã nguồn, bao gồm cả chức năng có API nhưng giao diện chưa tích hợp hoàn chỉnh.")

table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.LEFT
table.style = "Table Grid"
set_table_geometry(table, [800, 2200, 6360])

headers = ["STT", "Tên Actor", "Nhiệm vụ"]
for cell, value in zip(table.rows[0].cells, headers):
    set_cell_shading(cell, "1F4E79")
    write_cell(cell, value, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=(255, 255, 255))
set_repeat_table_header(table.rows[0])

actors = [
    ("1", "Quản trị viên", [
        "Xem và cập nhật hồ sơ quản trị viên.",
        "Theo dõi dashboard, thống kê, biểu đồ và báo cáo vận hành; xuất dữ liệu báo cáo trên giao diện.",
        "Quản lý tài khoản người dùng: tìm kiếm, xem, tạo, sửa, khóa/mở khóa và xử lý tài khoản theo quyền quản trị.",
        "Quản lý danh mục tour.",
        "Quản lý điểm đến.",
        "Quản lý tour: tạo, xem, sửa, xóa; quản lý hành trình, hình ảnh và quy tắc giá theo độ tuổi.",
        "Quản lý lịch khởi hành: tạo, cập nhật, theo dõi số chỗ và trạng thái lịch.",
        "Xem danh sách khách đã đặt theo từng lịch khởi hành.",
        "Quản lý booking: tra cứu, xem chi tiết và cập nhật trạng thái nghiệp vụ.",
        "Quản lý thanh toán: theo dõi, cập nhật trạng thái thanh toán và xử lý trạng thái hoàn tiền theo API.",
        "Quản lý hồ sơ hướng dẫn viên, bao gồm thông tin chuyên môn, ngôn ngữ, kinh nghiệm, chứng chỉ và ảnh đại diện.",
        "Phân công tự động hoặc chỉ định hướng dẫn viên cho lịch khởi hành; hủy phân công khi phù hợp.",
        "Xem, phê duyệt hoặc từ chối yêu cầu thay thế hướng dẫn viên.",
        "Xem, phê duyệt, từ chối hoặc cập nhật quyết định đối với đơn xin nghỉ của hướng dẫn viên.",
        "Quản lý nhân viên hỗ trợ: tạo, cập nhật, thống kê, xóa mềm, khôi phục, xóa vĩnh viễn và quản lý ảnh đại diện.",
        "Quản lý danh mục ngôn ngữ, cấp độ ngôn ngữ và chứng chỉ phục vụ hồ sơ hướng dẫn viên.",
        "Tìm kiếm, xem chi tiết và kiểm duyệt đánh giá tour (hiển thị, ẩn hoặc đánh dấu spam).",
        "Quản lý chiến dịch thông báo: chọn người nhận, lưu nháp, sửa, xóa, khôi phục, gửi, xem thông báo đã gửi và thu hồi.",
        "Đọc thông báo dành cho quản trị viên, đếm thông báo chưa đọc và đánh dấu đã đọc một hoặc tất cả.",
        "Cấu hình cài đặt hệ thống: hệ thống, bảo mật, thông báo, ngôn ngữ/múi giờ, thanh toán và sao lưu.",
        "Quản lý widget/banner: tạo, sửa, xóa và bật/tắt hiển thị.",
        "Quản lý sao lưu cơ sở dữ liệu: xem danh sách, tạo, tải xuống và xóa tệp sao lưu SQL.",
        "Quản lý loại dịch vụ qua API: tạo, xem, cập nhật và xóa mềm.",
    ]),
    ("2", "Hướng dẫn viên", [
        "Xem và cập nhật hồ sơ cá nhân/nghiệp vụ; thay đổi mật khẩu.",
        "Theo dõi dashboard hướng dẫn viên.",
        "Xem các tour và lịch khởi hành được phân công.",
        "Xem thông tin, danh sách khách tham gia của tour được phân công.",
        "Tạo và quản lý phiên điểm danh cho tour.",
        "Thực hiện check-in, check-out khách tham gia và cập nhật ghi chú điểm danh.",
        "Theo dõi và cập nhật các giai đoạn/tiến độ vận hành của tour.",
        "Gửi yêu cầu thay hướng dẫn viên khi không thể thực hiện tour.",
        "Theo dõi trạng thái yêu cầu thay hướng dẫn viên.",
        "Gửi đơn xin nghỉ kèm lý do và tệp minh chứng khi cần.",
        "Xem danh sách, thống kê và hủy đơn xin nghỉ còn ở trạng thái chờ xử lý.",
        "Xem đánh giá của khách hàng về bản thân và lịch sử tour đã thực hiện.",
        "Xem, đọc và theo dõi thông báo cá nhân.",
    ]),
    ("3", "Nhân viên hỗ trợ", [
        "Xem và cập nhật hồ sơ cá nhân; thay đổi mật khẩu.",
        "Theo dõi dashboard hỗ trợ và số lượng yêu cầu cần xử lý.",
        "Xem danh sách yêu cầu hỗ trợ; lọc theo từ khóa, trạng thái, danh mục và mức độ ưu tiên.",
        "Xem chi tiết yêu cầu, thông tin khách hàng và các tệp đính kèm.",
        "Tiếp nhận yêu cầu hỗ trợ, chuyển sang trạng thái đang xử lý hoặc hoàn tất xử lý.",
        "Theo dõi badge số yêu cầu đang chờ và đang xử lý.",
        "Sử dụng chatbot AI hỗ trợ trong không gian làm việc theo giao diện hiện có.",
        "Xem danh sách thông báo cá nhân và số thông báo chưa đọc.",
        "Xem chi tiết và đánh dấu thông báo cá nhân là đã đọc.",
        "Gửi thông báo đến toàn bộ quản trị viên khi cần phối hợp hoặc báo cáo sự việc.",
    ]),
    ("4", "Khách hàng", [
        "Đăng ký tài khoản khách hàng.",
        "Đăng nhập, duy trì phiên đăng nhập và đăng xuất.",
        "Yêu cầu OTP và đặt lại mật khẩu qua API khi quên mật khẩu.",
        "Xem trang chủ, danh mục, điểm đến và tour công khai.",
        "Tìm kiếm, lọc, sắp xếp và xem chi tiết tour cùng lịch khởi hành còn mở.",
        "Xem đánh giá tour công khai.",
        "Sử dụng trợ lý du lịch AI/chatbot để hỏi thông tin và gợi ý tour.",
        "Xem và cập nhật hồ sơ cá nhân, thay ảnh đại diện và đổi mật khẩu.",
        "Xem tổng quan tài khoản và lịch sử booking.",
        "Quản lý danh sách tour yêu thích: xem, thêm và xóa tour.",
        "Xem trước giá tour theo số lượng và độ tuổi hành khách.",
        "Tạo booking, khai báo thông tin liên hệ và danh sách hành khách.",
        "Thanh toán booking qua VNPAY.",
        "Tiếp tục thanh toán booking đang chờ hoặc hủy booking theo điều kiện hệ thống.",
        "Theo dõi trạng thái thanh toán VNPAY và trạng thái booking của chính mình.",
        "Tạo hoặc chỉnh sửa đánh giá tour khi đủ điều kiện.",
        "Đánh giá hướng dẫn viên sau chuyến đi khi đủ điều kiện; xem hồ sơ đánh giá liên quan.",
        "Gửi yêu cầu hỗ trợ kèm mô tả và tệp đính kèm; theo dõi phản hồi, bổ sung thông tin khi được yêu cầu.",
        "Xem, đọc và đánh dấu thông báo cá nhân.",
    ]),
]

for index, name, duty in actors:
    cells = table.add_row().cells
    write_cell(cells[0], index, align=WD_ALIGN_PARAGRAPH.CENTER)
    write_cell(cells[1], name, bold=True)
    write_bullets(cells[2], duty)

doc.add_paragraph().paragraph_format.space_after = Pt(0)
note = doc.add_paragraph()
note.paragraph_format.space_before = Pt(10)
note.paragraph_format.space_after = Pt(0)
note_run = note.add_run("Ghi chú: Danh sách chỉ bao gồm chức năng có bằng chứng trong mã nguồn tại thời điểm rà soát; các hạng mục chỉ có migration/seeder nhưng chưa có luồng API/UI hoàn chỉnh không được đưa vào.")
note_run.italic = True
note_run.font.size = Pt(9.5)
note_run.font.color.rgb = RGBColor(89, 89, 89)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
