from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path("Chi_tiet_CSDL_bo_sung_cac_bang_con_thieu.docx")


TABLES = [
    ("Attendance_activity_logs", [
        ("id", "bigint", "-", "v", "PK", "Khóa chính, tự tăng"),
        ("attendance_session_id", "bigint", "-", "v", "FK", "Phiên điểm danh liên quan"),
        ("booking_participant_id", "bigint", "-", "-", "FK", "Hành khách liên quan (nếu có)"),
        ("actor_id", "bigint", "-", "v", "FK", "Người thực hiện thao tác"),
        ("action", "varchar", "50", "v", "-", "Loại hành động điểm danh"),
        ("description", "varchar", "255", "v", "-", "Mô tả hành động"),
        ("metadata", "json", "-", "-", "-", "Dữ liệu bổ sung dạng JSON"),
        ("created_at", "timestamp", "-", "-", "-", "Thời gian tạo"),
        ("updated_at", "timestamp", "-", "-", "-", "Thời gian cập nhật"),
    ]),
    ("Attendance_session_photos", [
        ("id", "bigint", "-", "v", "PK", "Khóa chính, tự tăng"),
        ("attendance_session_id", "bigint", "-", "v", "FK", "Phiên điểm danh liên quan"),
        ("file_path", "varchar", "255", "v", "-", "Đường dẫn tệp ảnh"),
        ("original_name", "varchar", "255", "v", "-", "Tên tệp ảnh gốc"),
        ("uploaded_by", "bigint", "-", "v", "FK", "Người tải ảnh lên"),
        ("created_at", "timestamp", "-", "-", "-", "Thời gian tạo"),
        ("updated_at", "timestamp", "-", "-", "-", "Thời gian cập nhật"),
    ]),
    ("Booking_disruption_requests", [
        ("id", "bigint", "-", "v", "PK", "Khóa chính, tự tăng"),
        ("booking_id", "bigint", "-", "v", "FK", "Đơn đặt tour liên quan"),
        ("type", "enum", "-", "v", "-", "Hình thức xử lý: refund, retain, transfer"),
        ("status", "enum", "-", "v", "-", "Trạng thái: pending, approved, rejected"),
        ("reason", "text", "-", "-", "-", "Lý do khách hàng yêu cầu xử lý"),
        ("requested_tour_departure_id", "bigint", "-", "-", "FK", "Lịch khởi hành muốn chuyển sang"),
        ("admin_note", "text", "-", "-", "-", "Ghi chú của quản trị viên"),
        ("processed_by", "bigint", "-", "-", "FK", "Người xử lý yêu cầu"),
        ("processed_at", "timestamp", "-", "-", "-", "Thời gian xử lý"),
        ("created_at", "timestamp", "-", "-", "-", "Thời gian tạo"),
        ("updated_at", "timestamp", "-", "-", "-", "Thời gian cập nhật"),
    ]),
    ("Booking_information_change_histories", [
        ("id", "bigint", "-", "v", "PK", "Khóa chính, tự tăng"),
        ("booking_id", "bigint", "-", "v", "FK", "Đơn đặt tour được thay đổi"),
        ("changed_by", "bigint", "-", "v", "FK", "Người thực hiện thay đổi"),
        ("before", "json", "-", "-", "-", "Dữ liệu trước khi thay đổi"),
        ("after", "json", "-", "-", "-", "Dữ liệu sau khi thay đổi"),
        ("created_at", "timestamp", "-", "v", "-", "Thời gian ghi nhận thay đổi"),
    ]),
    ("Customer_presence_sessions", [
        ("id", "bigint", "-", "v", "PK", "Khóa chính, tự tăng"),
        ("user_id", "bigint", "-", "v", "FK", "Khách hàng tham gia phiên"),
        ("started_at", "timestamp", "-", "v", "-", "Thời gian bắt đầu phiên"),
        ("last_seen_at", "timestamp", "-", "v", "-", "Thời gian hoạt động gần nhất"),
        ("ended_at", "timestamp", "-", "-", "-", "Thời gian kết thúc phiên"),
        ("duration_seconds", "bigint", "-", "v", "-", "Thời lượng phiên (giây)"),
        ("ip_address", "varchar", "45", "-", "-", "Địa chỉ IP"),
        ("user_agent", "text", "-", "-", "-", "Thông tin trình duyệt/thiết bị"),
        ("created_at", "timestamp", "-", "-", "-", "Thời gian tạo"),
        ("updated_at", "timestamp", "-", "-", "-", "Thời gian cập nhật"),
    ]),
    ("Destination_place_activity_types", [
        ("id", "bigint", "-", "v", "PK", "Khóa chính, tự tăng"),
        ("destination_place_id", "bigint", "-", "v", "FK", "Địa điểm du lịch liên quan"),
        ("activity_type", "varchar", "40", "v", "UK", "Loại hoạt động tại địa điểm"),
        ("created_at", "timestamp", "-", "-", "-", "Thời gian tạo"),
        ("updated_at", "timestamp", "-", "-", "-", "Thời gian cập nhật"),
    ]),
    ("Destination_places", [
        ("id", "bigint", "-", "v", "PK", "Khóa chính, tự tăng"),
        ("province_id", "bigint", "-", "v", "FK", "Tỉnh/thành phố trực thuộc"),
        ("name", "varchar", "180", "v", "UK", "Tên địa điểm"),
        ("slug", "varchar", "220", "v", "UK", "Đường dẫn thân thiện"),
        ("district_name", "varchar", "150", "-", "-", "Tên quận/huyện lưu để tra cứu"),
        ("district_id", "bigint", "-", "-", "FK", "Quận/huyện liên quan"),
        ("address", "varchar", "500", "-", "-", "Địa chỉ chi tiết"),
        ("description", "text", "-", "-", "-", "Mô tả địa điểm"),
        ("thumbnail_url", "varchar", "500", "-", "-", "Ảnh đại diện"),
        ("status", "enum", "-", "v", "-", "Trạng thái: active, inactive"),
        ("deleted_at", "timestamp", "-", "-", "-", "Thời gian xóa mềm"),
        ("created_at", "timestamp", "-", "-", "-", "Thời gian tạo"),
        ("updated_at", "timestamp", "-", "-", "-", "Thời gian cập nhật"),
    ]),
    ("Districts", [
        ("id", "bigint", "-", "v", "PK", "Khóa chính, tự tăng"),
        ("province_id", "bigint", "-", "v", "FK", "Tỉnh/thành phố trực thuộc"),
        ("name", "varchar", "150", "v", "UK", "Tên quận/huyện"),
        ("code", "varchar", "20", "-", "-", "Mã đơn vị hành chính"),
        ("created_at", "timestamp", "-", "-", "-", "Thời gian tạo"),
        ("updated_at", "timestamp", "-", "-", "-", "Thời gian cập nhật"),
    ]),
    ("Faqs", [
        ("id", "bigint", "-", "v", "PK", "Khóa chính, tự tăng"),
        ("category", "varchar", "60", "v", "-", "Nhóm câu hỏi"),
        ("question", "varchar", "500", "v", "-", "Nội dung câu hỏi"),
        ("answer", "text", "-", "v", "-", "Nội dung trả lời"),
        ("keywords", "json", "-", "v", "-", "Từ khóa tìm kiếm dạng JSON"),
        ("sort_order", "smallint", "-", "v", "-", "Thứ tự hiển thị"),
        ("is_active", "tinyint", "1", "v", "-", "Trạng thái kích hoạt"),
        ("created_at", "timestamp", "-", "-", "-", "Thời gian tạo"),
        ("updated_at", "timestamp", "-", "-", "-", "Thời gian cập nhật"),
    ]),
    ("Guide_presence_sessions", [
        ("id", "bigint", "-", "v", "PK", "Khóa chính, tự tăng"),
        ("user_id", "bigint", "-", "v", "FK", "Tài khoản hướng dẫn viên"),
        ("started_at", "timestamp", "-", "v", "-", "Thời gian bắt đầu phiên"),
        ("last_seen_at", "timestamp", "-", "v", "-", "Thời gian hoạt động gần nhất"),
        ("ended_at", "timestamp", "-", "-", "-", "Thời gian kết thúc phiên"),
        ("duration_seconds", "bigint", "-", "v", "-", "Thời lượng phiên (giây)"),
        ("ip_address", "varchar", "45", "-", "-", "Địa chỉ IP"),
        ("user_agent", "text", "-", "-", "-", "Thông tin trình duyệt/thiết bị"),
        ("created_at", "timestamp", "-", "-", "-", "Thời gian tạo"),
        ("updated_at", "timestamp", "-", "-", "-", "Thời gian cập nhật"),
    ]),
    ("Provinces", [
        ("id", "bigint", "-", "v", "PK", "Khóa chính, tự tăng"),
        ("name", "varchar", "120", "v", "UK", "Tên tỉnh/thành phố"),
        ("code", "varchar", "20", "-", "UK", "Mã đơn vị hành chính"),
        ("created_at", "timestamp", "-", "-", "-", "Thời gian tạo"),
        ("updated_at", "timestamp", "-", "-", "-", "Thời gian cập nhật"),
    ]),
    ("Tour_activity_logs", [
        ("id", "bigint", "-", "v", "PK", "Khóa chính, tự tăng"),
        ("tour_id", "bigint", "-", "-", "FK", "Tour liên quan"),
        ("actor_id", "bigint", "-", "-", "FK", "Người thực hiện thao tác"),
        ("action", "varchar", "50", "v", "-", "Loại hành động"),
        ("tour_title", "varchar", "255", "v", "-", "Tên tour tại thời điểm ghi log"),
        ("description", "text", "-", "v", "-", "Mô tả chi tiết"),
        ("metadata", "json", "-", "-", "-", "Dữ liệu bổ sung dạng JSON"),
        ("created_at", "timestamp", "-", "-", "-", "Thời gian tạo"),
        ("updated_at", "timestamp", "-", "-", "-", "Thời gian cập nhật"),
    ]),
    ("Tour_departure_status_histories", [
        ("id", "bigint", "-", "v", "PK", "Khóa chính, tự tăng"),
        ("tour_departure_id", "bigint", "-", "v", "FK", "Lịch khởi hành liên quan"),
        ("old_status", "varchar", "50", "-", "-", "Trạng thái trước"),
        ("new_status", "varchar", "50", "v", "-", "Trạng thái mới"),
        ("reason", "varchar", "100", "-", "-", "Lý do thay đổi"),
        ("created_at", "timestamp", "-", "v", "-", "Thời gian ghi nhận"),
    ]),
    ("Tour_finalization_outbox", [
        ("id", "bigint", "-", "v", "PK", "Khóa chính, tự tăng"),
        ("tour_departure_id", "bigint", "-", "v", "FK", "Lịch khởi hành liên quan"),
        ("event_type", "varchar", "100", "v", "UK", "Loại sự kiện hoàn tất tour"),
        ("payload", "json", "-", "v", "-", "Dữ liệu sự kiện"),
        ("processed_at", "timestamp", "-", "-", "-", "Thời gian đã xử lý"),
        ("created_at", "timestamp", "-", "v", "-", "Thời gian tạo"),
    ]),
    ("Tour_refund_outbox", [
        ("id", "bigint", "-", "v", "PK", "Khóa chính, tự tăng"),
        ("booking_id", "bigint", "-", "v", "FK", "Đơn đặt tour liên quan"),
        ("refund_request_id", "bigint", "-", "v", "FK/UK", "Yêu cầu hoàn tiền liên quan"),
        ("payload", "json", "-", "-", "-", "Dữ liệu xử lý hoàn tiền"),
        ("processed_at", "timestamp", "-", "-", "-", "Thời gian đã xử lý"),
        ("created_at", "timestamp", "-", "v", "-", "Thời gian tạo"),
    ]),
]


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def margins(cell, top=80, start=100, bottom=80, end=100):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = OxmlElement(f"w:{tag}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


doc = Document()
section = doc.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width = Cm(29.7)
section.page_height = Cm(21.0)
section.top_margin = Cm(1.6)
section.bottom_margin = Cm(1.6)
section.left_margin = Cm(1.6)
section.right_margin = Cm(1.6)
section.header_distance = Cm(0.8)
section.footer_distance = Cm(0.8)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(4)
normal.paragraph_format.line_spacing = 1.15
for style_name in ("Title", "Heading 1"):
    style = styles[style_name]
    style.font.name = "Times New Roman"
    style.font.color.rgb = RGBColor(31, 78, 121)
styles["Title"].font.size = Pt(20)
styles["Heading 1"].font.size = Pt(13)
styles["Heading 1"].font.bold = True
styles["Heading 1"].paragraph_format.space_before = Pt(8)
styles["Heading 1"].paragraph_format.space_after = Pt(5)
styles["Heading 1"].paragraph_format.keep_with_next = True

title = doc.add_paragraph(style="Title")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.add_run("5.2. CHI TIẾT CƠ SỞ DỮ LIỆU - CÁC BẢNG BỔ SUNG")
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subtitle.add_run("Đối chiếu theo toàn bộ migration hiện có; tiếp nối sau Bảng 5.2.68")
r.italic = True
r.font.color.rgb = RGBColor(89, 89, 89)
subtitle.paragraph_format.space_after = Pt(12)

headers = ["No.", "Name", "Type", "Length", "Not Null", "Key", "Mô tả"]
widths = [Cm(0.9), Cm(4.5), Cm(2.1), Cm(1.6), Cm(1.8), Cm(1.5), Cm(11.0)]

for table_no, (table_name, columns) in enumerate(TABLES, start=69):
    heading = doc.add_paragraph(style="Heading 1")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.add_run(f"Bảng 5.2.{table_no}. Chi tiết bảng {table_name}")
    keep_with_next(heading)

    table = doc.add_table(rows=1, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    repeat_header(table.rows[0])
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.width = widths[idx]
        shade(cell, "D9EAF7")
        margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(9)

    for no, row_data in enumerate(columns, start=1):
        row = table.add_row()
        values = [str(no), *row_data]
        for idx, text in enumerate(values):
            cell = row.cells[idx]
            cell.width = widths[idx]
            margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx in (0, 2, 3, 4, 5) else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(9)
    if table_no != 83:
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(2)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer.add_run("Chi tiết cơ sở dữ liệu bổ sung từ migration")
fr.font.name = "Times New Roman"
fr.font.size = Pt(9)
fr.font.color.rgb = RGBColor(100, 100, 100)

core = doc.core_properties
core.title = "Chi tiết cơ sở dữ liệu - các bảng bổ sung"
core.subject = "Các bảng còn thiếu sau khi đối chiếu migration"
core.author = ""
doc.save(OUT)
print(OUT.resolve())
